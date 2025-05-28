import os
import re
import docx
import PyPDF2
import pandas as pd
from fpdf import FPDF
from docx import Document
from datetime import datetime

def extract_text_from_docx(file_path):
    try:
        doc = docx.Document(file_path)
        return '\n'.join([p.text for p in doc.paragraphs])
    except Exception:
        return ""

def extract_text_from_pdf(file_path):
    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            return "\n".join([page.extract_text() or "" for page in reader.pages])
    except Exception:
        return ""

def extract_keywords(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".txt":
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read().split(',')
    elif ext == ".docx":
        content = extract_text_from_docx(file_path)
        return content.split(',')
    return []

def process_resume(resume_path, keywords):
    matched_keywords = []
    score = 0

    ext = os.path.splitext(resume_path)[1].lower()

    if ext == '.pdf':
        content = extract_text_from_pdf(resume_path).lower()
    elif ext == '.docx':
        content = extract_text_from_docx(resume_path).lower()
    elif ext == '.txt':
        with open(resume_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read().lower()
    else:
        content = ""  # unsupported type

    # Match whole keywords only
    for kw in keywords:
        if re.search(r'\b' + re.escape(kw.lower()) + r'\b', content):
            matched_keywords.append(kw)
            score += 1

    return score, matched_keywords

def export_results(results, format='csv'):
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"results_{timestamp}.{format}"
    filepath = os.path.join('results', filename)
    df = pd.DataFrame(results)

    if format == 'csv':
        df.to_csv(filepath, index=False)
    elif format == 'txt':
        with open(filepath, 'w', encoding='utf-8') as f:
            for _, row in df.iterrows():
                f.write(f"{row['filename']}: {row['match_score']} matches\n")
    elif format == 'docx':
        document = Document()
        document.add_heading("CV Screening Results", level=1)
        for _, row in df.iterrows():
            document.add_paragraph(f"{row['filename']}: {row['match_score']} matches")
        document.save(filepath)
    elif format == 'pdf':
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, "CV Screening Results", ln=True)
        for _, row in df.iterrows():
            pdf.cell(200, 10, f"{row['filename']}: {row['match_score']} matches", ln=True)
        pdf.output(filepath)

    return filepath
